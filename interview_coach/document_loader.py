"""Load and process documents for knowledge base."""

import json
from pathlib import Path
from typing import List, Optional


class Document:
    """Simple document representation."""
    def __init__(self, content: str, source: str, doc_type: str = "text"):
        self.content = content
        self.source = source
        self.doc_type = doc_type


class DocumentLoader:
    """Load documents from various formats. Configurable file type support."""

    # Default supported formats
    DEFAULT_FORMATS = {
        "txt": "plain_text",
        "md": "markdown",
        "html": "html",
        "pdf": "pdf",
        "docx": "docx",
    }

    def __init__(self, documents_path: Path, career_history_file: Path = None,
                 supported_formats: Optional[dict] = None):
        self.documents_path = documents_path
        self.career_history_file = career_history_file
        self.documents = []
        # Allow override of supported formats
        self.supported_formats = supported_formats or self.DEFAULT_FORMATS
        self._loaders = {
            "plain_text": self._load_plain_text,
            "markdown": self._load_plain_text,  # markdown is plain text
            "html": self._load_html,
            "pdf": self._load_pdf,
            "docx": self._load_docx,
        }

    def load_all(self) -> List[Document]:
        """Load all documents and career history."""
        self.documents = []

        # Load all supported file formats
        if self.documents_path.exists():
            for ext, format_type in self.supported_formats.items():
                for file in self.documents_path.glob(f"**/*.{ext}"):
                    loader = self._loaders.get(format_type)
                    if loader:
                        loader(file)

        # Load career history
        if self.career_history_file and self.career_history_file.exists():
            self._load_career_history()

        return self.documents

    def _load_plain_text(self, filepath: Path):
        """Load plain text or markdown file."""
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            self.documents.append(Document(content, str(filepath)))
        except Exception as e:
            print(f"Error loading {filepath}: {e}")

    def _load_html(self, filepath: Path):
        """Load HTML file (basic extraction)."""
        try:
            # Optional: try beautiful soup if available, else basic regex
            try:
                from html.parser import HTMLParser

                class TextExtractor(HTMLParser):
                    def __init__(self):
                        super().__init__()
                        self.text = []
                    def handle_data(self, data):
                        if data.strip():
                            self.text.append(data)

                with open(filepath, 'r', encoding='utf-8') as f:
                    html = f.read()

                parser = TextExtractor()
                parser.feed(html)
                content = "\n".join(parser.text)

                self.documents.append(Document(content, str(filepath)))
            except Exception:
                # Fallback: just read as text
                with open(filepath, 'r', encoding='utf-8') as f:
                    content = f.read()
                self.documents.append(Document(content, str(filepath)))
        except Exception as e:
            print(f"Error loading HTML {filepath}: {e}")

    def _load_pdf(self, filepath: Path):
        """Load PDF file if pdfplumber or pypdf is available."""
        try:
            try:
                import pdfplumber
                with pdfplumber.open(filepath) as pdf:
                    text = "\n".join(page.extract_text() or "" for page in pdf.pages)
                self.documents.append(Document(text, str(filepath)))
            except ImportError:
                try:
                    from PyPDF2 import PdfReader
                    reader = PdfReader(filepath)
                    text = "\n".join(page.extract_text() or "" for page in reader.pages)
                    self.documents.append(Document(text, str(filepath)))
                except ImportError:
                    print(f"Warning: No PDF library available (pdfplumber or PyPDF2). Skipping {filepath}")
        except Exception as e:
            print(f"Error loading PDF {filepath}: {e}")

    def _load_docx(self, filepath: Path):
        """Load DOCX file if python-docx is available."""
        try:
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(filepath)
                text = "\n".join(para.text for para in doc.paragraphs)
                self.documents.append(Document(text, str(filepath)))
            except ImportError:
                print(f"Warning: python-docx not available. Install with: pip install python-docx. Skipping {filepath}")
        except Exception as e:
            print(f"Error loading DOCX {filepath}: {e}")

    def _load_career_history(self):
        """Load and parse career_history.json."""
        try:
            with open(self.career_history_file, 'r') as f:
                career_data = json.load(f)

            # Convert career data to searchable text
            text_parts = []
            if "summary" in career_data:
                text_parts.append(f"Career Summary: {career_data['summary']}")
            if "skills" in career_data:
                text_parts.append(f"Skills: {', '.join(career_data['skills'])}")
            if "experience" in career_data:
                for exp in career_data["experience"]:
                    text_parts.append(f"Experience: {exp.get('role')} at {exp.get('company')} ({exp.get('duration')})")
                    if "achievements" in exp:
                        text_parts.append(f"Achievements: {', '.join(exp['achievements'])}")

            content = "\n".join(text_parts)
            self.documents.append(Document(content, str(self.career_history_file), "career"))
        except Exception as e:
            print(f"Error loading career history: {e}")
